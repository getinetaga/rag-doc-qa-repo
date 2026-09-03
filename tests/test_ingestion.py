"""Tests for `app.ingestion` helpers.

These tests verify that `extract_text` correctly handles plain text and
DOCX files, that encoding fallbacks are honored, and that unsupported
extensions raise `ValueError`.
"""

import base64
from pathlib import Path
import tempfile
import pytest

from app import ingestion
from app.ingestion import (
    _encode_sharing_url,
    _get_graph_token,
    _reset_graph_token_cache,
    extract_google_doc_text,
    extract_sharepoint_text,
    extract_text,
)


def test_extract_txt(tmp_path: Path):
    p = tmp_path / "sample.txt"
    content = "Hello world\nThis is a test."
    p.write_text(content, encoding="utf-8")

    out = extract_text(str(p))
    assert "Hello world" in out
    assert "This is a test." in out


def test_extract_docx(tmp_path: Path):
    # Create a small DOCX file using python-docx
    try:
        import docx
    except Exception:
        pytest.skip("python-docx not available")

    p = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("Paragraph one")
    doc.add_paragraph("Paragraph two")
    doc.save(str(p))

    out = extract_text(str(p))
    assert "Paragraph one" in out
    assert "Paragraph two" in out


def test_extract_unsupported_extension(tmp_path: Path):
    p = tmp_path / "data.bin"
    p.write_bytes(b"\x00\x01\x02")

    try:
        extract_text(str(p))
    except ValueError as e:
        assert "Unsupported file type" in str(e)
    else:
        raise AssertionError("Expected ValueError for unsupported extension")


def test_extract_pdf_with_mocked_pdfplumber(monkeypatch, tmp_path: Path):
    p = tmp_path / "sample.pdf"
    p.write_bytes(b"%PDF-1.4\n%fake")

    class _Page:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class _PDF:
        def __init__(self):
            self.pages = [_Page("Page 1 text"), _Page(None), _Page("Page 3 text")]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

    def fake_open(path):
        assert str(path).endswith("sample.pdf")
        return _PDF()

    monkeypatch.setattr("app.ingestion.pdfplumber.open", fake_open)

    out = extract_text(str(p))
    assert "Page 1 text" in out
    assert "Page 3 text" in out
    assert "None" not in out


def test_extract_txt_latin1_fallback(tmp_path: Path):
    p = tmp_path / "latin1.txt"
    p.write_bytes("cafe\xe9".encode("latin-1"))

    out = extract_text(str(p))
    assert out == "cafeé"


def test_extract_missing_file_raises_file_not_found(tmp_path: Path):
    missing = tmp_path / "does_not_exist.txt"
    with pytest.raises(FileNotFoundError):
        extract_text(str(missing))


def test_extract_google_doc_text(monkeypatch):
    class FakeResponse:
        text = "Google Docs content"

        def raise_for_status(self):
            return None

    def fake_get(url, timeout):
        assert "export?format=txt" in url
        assert timeout == 20
        return FakeResponse()

    monkeypatch.setattr("app.ingestion.requests.get", fake_get)

    out = extract_google_doc_text("https://docs.google.com/document/d/abc123/edit")
    assert out == "Google Docs content"


def test_extract_google_doc_text_invalid_url():
    with pytest.raises(ValueError, match="Invalid Google Docs URL"):
        extract_google_doc_text("https://example.com/not-a-google-doc")


# ---------------------------------------------------------------------------
# SharePoint / Microsoft Graph ingestion
# ---------------------------------------------------------------------------


class _FakeResponse:
    def __init__(self, *, json_data=None, content=b""):
        self._json = json_data or {}
        self.content = content

    def raise_for_status(self):
        return None

    def json(self):
        return self._json


@pytest.fixture
def graph_credentials(monkeypatch):
    """Configure app-only Graph credentials and clear the token cache."""

    monkeypatch.setattr(ingestion.config, "SHAREPOINT_TENANT_ID", "tenant-abc")
    monkeypatch.setattr(ingestion.config, "SHAREPOINT_CLIENT_ID", "client-abc")
    monkeypatch.setattr(ingestion.config, "SHAREPOINT_CLIENT_SECRET", "secret-abc")
    monkeypatch.setattr(ingestion.config, "GRAPH_BASE_URL", "https://graph.microsoft.com/v1.0")
    monkeypatch.setattr(ingestion.config, "GRAPH_AUTHORITY", "https://login.microsoftonline.com")
    _reset_graph_token_cache()
    yield
    _reset_graph_token_cache()


def test_encode_sharing_url_matches_graph_shares_format():
    encoded = _encode_sharing_url("https://contoso.sharepoint.com/sites/x/Doc.docx")

    assert encoded.startswith("u!")
    assert "=" not in encoded
    decoded = base64.urlsafe_b64decode(encoded[2:] + "===").decode("utf-8")
    assert decoded == "https://contoso.sharepoint.com/sites/x/Doc.docx"


def test_get_graph_token_requires_credentials(monkeypatch):
    monkeypatch.setattr(ingestion.config, "SHAREPOINT_TENANT_ID", None)
    monkeypatch.setattr(ingestion.config, "SHAREPOINT_CLIENT_ID", None)
    monkeypatch.setattr(ingestion.config, "SHAREPOINT_CLIENT_SECRET", None)
    _reset_graph_token_cache()

    with pytest.raises(ValueError, match="SHAREPOINT_TENANT_ID"):
        _get_graph_token()


def test_get_graph_token_caches_token(monkeypatch, graph_credentials):
    calls = {"count": 0}

    def fake_post(url, data=None, timeout=None):
        calls["count"] += 1
        assert "oauth2/v2.0/token" in url
        assert data["grant_type"] == "client_credentials"
        return _FakeResponse(json_data={"access_token": "tok-123", "expires_in": 3600})

    monkeypatch.setattr(ingestion.requests, "post", fake_post)

    assert _get_graph_token() == "tok-123"
    assert _get_graph_token() == "tok-123"
    assert calls["count"] == 1


def test_extract_sharepoint_text_downloads_and_extracts(monkeypatch, graph_credentials):
    monkeypatch.setattr(
        ingestion.requests,
        "post",
        lambda url, data=None, timeout=None: _FakeResponse(
            json_data={"access_token": "tok-123", "expires_in": 3600}
        ),
    )

    def fake_get(url, headers=None, timeout=None, allow_redirects=False):
        assert headers["Authorization"] == "Bearer tok-123"
        if url.endswith("/driveItem/content"):
            return _FakeResponse(content=b"SharePoint policy body text.")
        assert url.endswith("/driveItem")
        return _FakeResponse(json_data={"name": "policy.txt", "size": 27})

    monkeypatch.setattr(ingestion.requests, "get", fake_get)

    out = extract_sharepoint_text(
        "https://contoso.sharepoint.com/sites/hr/Shared%20Documents/policy.txt"
    )
    assert out == "SharePoint policy body text."


def test_extract_sharepoint_text_rejects_unsupported_type(monkeypatch, graph_credentials):
    monkeypatch.setattr(
        ingestion.requests,
        "post",
        lambda url, data=None, timeout=None: _FakeResponse(
            json_data={"access_token": "tok-123", "expires_in": 3600}
        ),
    )
    monkeypatch.setattr(
        ingestion.requests,
        "get",
        lambda url, headers=None, timeout=None, allow_redirects=False: _FakeResponse(
            json_data={"name": "archive.zip", "size": 10}
        ),
    )

    with pytest.raises(ValueError, match="Unsupported SharePoint file type"):
        extract_sharepoint_text("https://contoso.sharepoint.com/sites/x/archive.zip")


def test_extract_sharepoint_text_requires_a_locator(graph_credentials):
    with pytest.raises(ValueError, match="Provide a sharepoint_url"):
        extract_sharepoint_text()